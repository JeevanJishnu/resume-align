from docx import Document
from .template_models import TemplateSchema, SectionType
import logging
import re

logger = logging.getLogger(__name__)

def _normalize(text):
    """Normalize text for fuzzy matching."""
    if not text: return ""
    return re.sub(r'[^a-zA-Z0-9\s]', '', text.lower()).strip()

def _find_placeholders_in_cell(cell, header_text, all_headers=None):
    """
    Find all placeholder paragraphs after a specific header within a cell.
    """
    placeholders = []
    found_header = False
    
    h_norm = _normalize(header_text)
    other_headers_norm = []
    if all_headers:
        for h in all_headers:
            hn = _normalize(h)
            if hn and hn != h_norm:
                other_headers_norm.append(hn)

    for p in cell.paragraphs:
        txt_raw = p.text
        txt_norm = _normalize(txt_raw)
        
        if not found_header:
            if h_norm in txt_norm and len(txt_norm) < len(h_norm) + 20:
                found_header = True
                # If the header itself is a placeholder, include it
                if "[FILL HERE]" in txt_raw or "[CONTENT HERE]" in txt_raw:
                    placeholders.append(p)
            continue
            
        if other_headers_norm and txt_norm:
            if any(oh == txt_norm or (oh in txt_norm and len(txt_norm) < len(oh) + 10) for oh in other_headers_norm):
                break
        
        if "[FILL HERE]" in txt_raw or "[CONTENT HERE]" in txt_raw:
            placeholders.append(p)
            
    return placeholders

def fill_template(schema: TemplateSchema, data: dict, output_path: str):
    """Main entry point to fill CV data into a template schema."""
    try:
        doc = Document(schema.template_file)
        all_headers = [sect.header_text for sect in schema.sections]
        
        # Track consumed data to implement the "already pasted" logic
        consumed_content = {
            SectionType.SUMMARY: False,
            SectionType.SKILLS: False,
            SectionType.WORK_EXPERIENCE: False,
            SectionType.PROJECTS: False,
            SectionType.EDUCATION: False,
            SectionType.CERTIFICATIONS: False,
            SectionType.FULL_NAME: False,
            SectionType.TOOLS: False
        }
        
        # Track paragraphs that have been filled to ensure instance uniqueness
        filled_xml_ids = set()

        # Global Personal Info Replacement
        personal_map = {
            "{{full_name}}": data.get("full_name", "Applicant"),
            "{{email}}": data.get("email", "N/A"),
            "{{phone}}": data.get("phone", "N/A"),
            "{{linkedin}}": data.get("linkedin", "N/A"),
            "[NAME]": data.get("full_name", "Applicant"),
            "[fill Name here]": data.get("full_name", "Applicant")
        }
        
        # 1. Direct Name Replacement (Top of document logic)
        name_filled = False
        for p in doc.paragraphs:
            text_lower = p.text.lower()
            for k, v in personal_map.items():
                if k.lower() in text_lower:
                    pattern = re.compile(re.escape(k), re.IGNORECASE)
                    p.text = pattern.sub(v, p.text)
                    if k.lower() in ["{{full_name}}", "[name]", "[fill name here]"]: name_filled = True
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text_lower = p.text.lower()
                        for k, v in personal_map.items():
                            if k.lower() in text_lower:
                                pattern = re.compile(re.escape(k), re.IGNORECASE)
                                p.text = pattern.sub(v, p.text)
                                if k.lower() in ["{{full_name}}", "[name]", "[fill name here]"]: name_filled = True

        for sect in schema.sections:
            # Skip full_name section if we already filled it via global replace
            if sect.section_type == SectionType.FULL_NAME:
                continue

            text_data = ""
            list_data = []
            
            # Check mapping strategy for skills/tools
            has_skills_box = any(s.section_type == SectionType.SKILLS for s in schema.sections)
            has_tools_box = any(s.section_type == SectionType.TOOLS for s in schema.sections)
            
            # Check if this section type has already been consumed
            if consumed_content.get(sect.section_type) and sect.section_type in [SectionType.SUMMARY, SectionType.SKILLS, SectionType.TOOLS]:
                text_data = "" 
                list_data = []
            else:
                if sect.section_type == SectionType.SUMMARY:
                    text_data = data.get("summary", "").strip()
                    # Remove leading bullet if present
                    text_data = re.sub(r'^[•▪\-*]\s*', '', text_data)
                elif sect.section_type == SectionType.SKILLS:
                    if has_tools_box:
                        # Split mode: Skills box gets only tech skills
                        list_data = data.get("skills", [])
                    else:
                        # Unified mode: Skills box gets both
                        list_data = data.get("skills", []) + data.get("tools", [])
                    text_data = ", ".join(list_data)
                elif sect.section_type == SectionType.TOOLS:
                    # Tools box always gets tools
                    list_data = data.get("tools", [])
                    text_data = ", ".join(list_data)
                elif sect.section_type == SectionType.WORK_EXPERIENCE:
                    list_data = data.get("work_experience", [])
                elif sect.section_type == SectionType.PROJECTS:
                    list_data = data.get("projects", [])
                elif sect.section_type == SectionType.EDUCATION:
                    list_data = data.get("education", [])
                elif sect.section_type == SectionType.CERTIFICATIONS:
                    list_data = data.get("certifications", [])

            # Process filling
            filled_anything = False
            if sect.table_idx is not None and sect.table_idx < len(doc.tables):
                table = doc.tables[sect.table_idx]
                for row in table.rows:
                    for cell in row.cells:
                        targets = _find_placeholders_in_cell(cell, sect.header_text, all_headers)
                        if targets:
                            if sect.section_type in [SectionType.SUMMARY, SectionType.SKILLS, "tools"]:
                                if text_data:
                                    targets[0].text = text_data
                                    fmt = targets[0].paragraph_format
                                    fmt.space_after = 0
                                    fmt.space_before = 0
                                    fmt.line_spacing = 1.0
                                    filled_anything = True
                                else:
                                    # Physically remove if possible
                                    p_elem = targets[0]._element
                                    p_elem.getparent().remove(p_elem)
                                
                                # Remove all subsequent placeholders for this section
                                for rem in targets[1:]:
                                    p_elem = rem._element
                                    p_elem.getparent().remove(p_elem)
                            else:
                                for target in targets:
                                    if list_data:
                                        _fill_list_to_paragraph(target, sect.section_type, list_data)
                                        filled_anything = True
                                    elif text_data:
                                        target.text = text_data
                                        filled_anything = True
                                    else:
                                        # Physically remove unused placeholders
                                        p_elem = target._element
                                        p_elem.getparent().remove(p_elem)
            
            elif sect.start_paragraph_idx is not None:
                idx = sect.start_paragraph_idx + 1
                if idx < len(doc.paragraphs):
                    target = doc.paragraphs[idx]
                    if sect.section_type in [SectionType.WORK_EXPERIENCE, SectionType.PROJECTS, SectionType.EDUCATION] and list_data:
                        _fill_list_to_paragraph(target, sect.section_type, list_data)
                        filled_anything = True
                    elif text_data:
                        target.text = text_data
                        filled_anything = True
                    else:
                        p_elem = target._element
                        p_elem.getparent().remove(p_elem)

            if filled_anything:
                if sect.section_type in [SectionType.SUMMARY, SectionType.SKILLS, SectionType.TOOLS]:
                    consumed_content[sect.section_type] = True
                # For SKILLS, we allow filling different boxes (Key Skills, Other Skills) 
                # but we could mark as consumed if we want to stop after one.
                # Let's say we only consume if they have the same header.
                # Actually, let's keep it simple: allow all skills boxes.

        doc.save(output_path)
    except Exception as e:
        logger.error(f"Error in fill_template: {e}")

def _fill_list_to_paragraph(target_p, section_type, items):
    """Universal bullet-fill with strict formatting and cleanup."""
    # Store parent to remove target_p later
    p_element = target_p._element
    parent = p_element.getparent()
    
    noise_re = re.compile(r'^[•▪\-\*▪\x00-\x1f\x7f-\x9f\s\t]+')
    
    # Track used items to avoid semantic duplication in mirrored cells
    for item in reversed(items):
        p = target_p.insert_paragraph_before()
        p.paragraph_format.space_after = 0
        if section_type == SectionType.WORK_EXPERIENCE:
            role = noise_re.sub('', item.get('role', 'Role')).strip()
            comp = noise_re.sub('', item.get('company', 'Company')).strip()
            run = p.add_run(f"• {role} at {comp} ({item.get('duration', 'N/A')})")
            run.bold = True
            desc = noise_re.sub('', item.get('responsibilities', '')).strip()
            if desc: 
                d_p = target_p.insert_paragraph_before(desc + "\n")
                d_p.paragraph_format.space_after = 0
        elif section_type == SectionType.PROJECTS:
            title = noise_re.sub('', item.get('title', 'Project')).strip()
            run = p.add_run(f"• {title} ({item.get('duration', 'N/A')})")
            run.bold = True
            role = noise_re.sub('', item.get('role', 'N/A')).strip()
            tech = noise_re.sub('', item.get('tech', 'N/A')).strip()
            desc = f"Role: {role}\nTech Stack: {tech}\n"
            d_p = target_p.insert_paragraph_before(desc)
            d_p.paragraph_format.space_after = 0
        elif section_type == SectionType.EDUCATION:
            deg = noise_re.sub('', item.get('degree', 'Degree')).strip()
            inst = noise_re.sub('', item.get('institution', 'N/A')).strip()
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(f"• {deg} from {inst} ({item.get('duration', 'N/A')})")
            run.bold = True

    # Remove the original [FILL HERE] paragraph to eliminate the ghost line
    parent.remove(p_element)
