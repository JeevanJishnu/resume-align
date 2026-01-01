import re
from docx import Document
from .template_models import TemplateSchema, SectionType
from fuzzywuzzy import fuzz
import logging

logger = logging.getLogger("NuclearCleaner")
logger.setLevel(logging.INFO)

def normalize_text(text):
    if not text: return ""
    return re.sub(r'[\[\]\(\):]', '', text).strip().lower()

def clean_template_content(doc: Document, schema: TemplateSchema) -> Document:
    """
    STABLE SINGLETON CLEANER:
    - Unified normalization.
    - Recursive table protection.
    - Explicitly clears Name and replaces with [fill Name here].
    """
    claimed_headers = set() 
    
    def process_sequence(paragraphs):
        active_type = None
        has_slot = False
        to_del = []
        
        for p in paragraphs:
            txt = p.text.strip()
            norm_p = normalize_text(txt)
            
            matched_sect = None
            if txt:
                # Find which specific schema section this belongs to
                for sect in schema.sections:
                    h_norm = normalize_text(sect.header_text)
                    thr = 65 if sect.section_type == SectionType.FULL_NAME else 80
                    if fuzz.ratio(norm_p, h_norm) > thr:
                        matched_sect = sect
                        break
            
            if matched_sect:
                sect_id = f"{matched_sect.section_type}_{normalize_text(matched_sect.header_text)}"
                
                if matched_sect.section_type == SectionType.FULL_NAME:
                    if sect_id not in claimed_headers:
                        claimed_headers.add(sect_id)
                        p.clear()
                        p.add_run("[fill Name here]") # Explicit Placeholder
                        active_type = None
                        continue
                    else:
                        to_del.append(p)
                        continue

                if sect_id not in claimed_headers:
                    claimed_headers.add(sect_id)
                    active_type = matched_sect.section_type
                    has_slot = False
                    continue
                else:
                    # Duplicate header instance? Treat as ghost to clean it
                    to_del.append(p)
                    active_type = "GHOST"
                    continue
            
            if active_type == "GHOST":
                 to_del.append(p)
                 continue

            if active_type:
                if not has_slot:
                    p.clear()
                    p.add_run("[FILL HERE]")
                    has_slot = True
                else:
                    to_del.append(p)
            else:
                # If we are in NO active section but there's text, it's either junk or contact info
                # that wasn't in the schema. In a cleaned template, we want it gone.
                if txt: to_del.append(p)

        for p_rem in to_del:
            try:
                el = p_rem._element
                el.getparent().remove(el)
            except: pass

    # Run on Body
    process_sequence(doc.paragraphs)
    
    # Run on Tables
    seen_ids = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                eid = id(cell._element)
                if eid not in seen_ids:
                    process_sequence(cell.paragraphs)
                    seen_ids.add(eid)
                    # Recurse for nested
                    for nt in cell.tables:
                         for nr in nt.rows:
                             for nc in nr.cells:
                                 nid = id(nc._element)
                                 if nid not in seen_ids:
                                     process_sequence(nc.paragraphs)
                                     seen_ids.add(nid)
    return doc
