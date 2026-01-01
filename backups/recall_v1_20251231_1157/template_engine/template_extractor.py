import re
import logging
from docx import Document
from docx.oxml.ns import qn
from fuzzywuzzy import fuzz
from .template_models import TemplateSchema, TemplateSection, SectionType

logger = logging.getLogger(__name__)

# Configurable Synonyms
SECTION_SYNONYMS = {
    SectionType.SUMMARY: ["Summary of Qualifications", "Professional Summary", "Profile Summary", "Career Snapshot", "Technical Summary", "Professional Profile", "Career Summary", "Executive Summary"],
    SectionType.WORK_EXPERIENCE: ["Work Experience", "Professional Experience", "Employment History", "Experience Summary", "Career History", "Work Exp", "Experience", "Internships", "Industrial Training"],
    SectionType.SKILLS: ["Skills", "Technical Skills", "Core Competencies", "Key Skills", "Skill Set", "Areas of Expertise", "Key Skills and Knowledge", "Tech Stack", "Technical Proficiencies", "Core Skills", "IT Skills", "Professional Skills", "Software Skills", "Programming Skills", "Technical Competencies", "Other Skills"],
    SectionType.TOOLS: ["Tools", "Tools & Technologies", "Design Tools", "System Architecture Tools", "Technical Tools"],
    SectionType.PROJECTS: ["Projects", "Key Projects", "Selected Projects", "Project Experience", "Project Portfolio", "Assignments", "Project Details"],
    SectionType.EDUCATION: ["Education", "Academic Background", "Educational Qualifications", "Academic Profile", "Academic Qualification", "Academics"],
    SectionType.CERTIFICATIONS: ["Certifications", "Professional Certifications", "Licenses & Certifications", "Credentials", "Courses", "Awards", "Training and Certifications"]
}

def mark_protected_xml(element):
    element._element.set(qn('w:protected'), 'true')

def is_protected_xml(element):
    return element._element.get(qn('w:protected')) == 'true'

def iterate_doc_content(doc: Document):
    parent_elm = doc._body._element
    t_idx = 0
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            para = [p for p in doc.paragraphs if p._element == child]
            if para: yield ('para', para[0])
        elif child.tag.endswith('tbl'):
            table = doc.tables[t_idx]
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        if p.text.strip():
                            setattr(p, '_parent_table_idx', t_idx)
                            setattr(p, '_parent_cell_col', c_idx)
                            yield ('table_cell', p)
            t_idx += 1

def is_section_header(element) -> tuple[SectionType, float]:
    text_clean = element.text.strip()
    if not text_clean: return SectionType.UNKNOWN, 0.0
    if text_clean.startswith(('-', '•', '▪', '*', '➢')): return SectionType.UNKNOWN, 0.0
    skip_prefixes = ["Role:", "Client:", "Tools:", "Environment:", "Technologies:", "Duration:"]
    if any(text_clean.lower().startswith(p.lower()) for p in skip_prefixes): return SectionType.UNKNOWN, 0.0

    text_upper = text_clean.upper()
    is_bold = any(run.bold for run in element.runs)
    word_count = len(text_clean.split())
    best_type = SectionType.UNKNOWN
    best_score = 0
    
    if word_count > 6: return SectionType.UNKNOWN, 0.0

    for stype, synonyms in SECTION_SYNONYMS.items():
        for syn in synonyms:
            score = fuzz.token_set_ratio(text_upper, syn.upper())
            if is_bold: score += 30
            if stype in [SectionType.SKILLS, SectionType.TOOLS]:
                col_idx = getattr(element, '_parent_cell_col', -1)
                if col_idx in [1, 2]: score += 25
            if len(text_clean) > len(syn) + 15 and stype != SectionType.SUMMARY: score -= 30
            if text_clean.lower() == syn.lower(): score += 20
            if score > best_score:
                best_score = score
                best_type = stype

    if best_score >= 85: return best_type, min(1.0, best_score/100.0)
    return SectionType.UNKNOWN, 0.0

def detect_personal_info(doc: Document) -> list:
    candidates = []
    content = list(iterate_doc_content(doc))[:20]
    max_size = 0
    name_cand = None
    
    for i, (etype, elem) in enumerate(content):
        # Calculate max font size in this element
        sizes = [r.font.size.pt for r in elem.runs if r.font.size]
        word_count = len(elem.text.strip().split())
        is_very_top = i < 2
        
        if sizes:
            m = max(sizes)
            # Find all spans with characters
            name_parts = [r.text.strip() for r in elem.runs if r.text.strip()]
            full_name = " ".join(dict.fromkeys(name_parts))
            
            # GUARD: A name should not be more than 4 words or 40 chars
            if len(full_name.split()) <= 4 and len(full_name) < 40 and word_count <= 4:
                if m > max_size or (is_very_top and not name_cand):
                    max_size = m
                    name_cand = (i, etype, elem)
            elif is_very_top and word_count <= 4 and not name_cand:
                name_cand = (i, etype, elem)
        elif is_very_top and word_count <= 4 and not name_cand:
            name_cand = (i, etype, elem)
    
    if name_cand:
        idx, etype, elem = name_cand
        candidates.append({
            'index': idx,
            'type': SectionType.FULL_NAME,
            'confidence': 1.0,
            'text': elem.text.strip(),
            'elem_type': etype,
            'table_idx': getattr(elem, '_parent_table_idx', None),
            'xml_id': id(elem._element)
        })

    for i, (etype, elem) in enumerate(content):
        txt = elem.text.lower()
        if "@" in txt or re.search(r'\d{5,}', txt) or "linkedin" in txt:
            if name_cand and i == name_cand[0]: continue
            candidates.append({
                'index': i, 'type': SectionType.CONTACT_INFO, 'confidence': 0.9,
                'text': elem.text.strip(), 'elem_type': etype,
                'table_idx': getattr(elem, '_parent_table_idx', None), 'xml_id': id(elem._element)
            })
    return candidates

def _denoise_candidates(candidates: list) -> list:
    if not candidates: return []
    text_freq = {}
    for cand in candidates:
        text = cand['text'].lower().strip()
        text_freq[text] = text_freq.get(text, 0) + 1
    freq_filtered = [c for c in candidates if text_freq[c['text'].lower().strip()] <= 2]
    length_filtered = [c for c in freq_filtered if len(c['text']) >= 5 or c['confidence'] > 110]
    colon_filtered = [c for c in length_filtered if not c['text'].strip().endswith(':') or c['confidence'] > 95]
    blacklist = ["role", "client", "duration", "company name", "c2cz", "sketch", "user experience design", "logo design"]
    final_blacklist_filtered = [c for c in colon_filtered if not any(b in c['text'].lower() for b in blacklist)]

    type_to_best = {}
    for cand in final_blacklist_filtered:
        stype = cand['type']
        if stype == SectionType.UNKNOWN: continue
        if stype == SectionType.SKILLS:
            key = (stype, cand['text'].lower().strip())
            if key not in type_to_best: type_to_best[key] = cand
            continue
        if stype not in type_to_best or cand['confidence'] > type_to_best[stype]['confidence']:
            type_to_best[stype] = cand
    unique_sections = list(type_to_best.values())
    final_sections = sorted(unique_sections, key=lambda x: x['index'])
    logger.info(f"🧹 DENOISE STATS: {len(candidates)} -> {len(final_sections)} (SINGLETON ENFORCED)")
    return final_sections

def extract_template_schema(template_path: str, template_name: str, doc: Document = None) -> TemplateSchema:
    if doc is None: doc = Document(template_path)
    personal_candidates = detect_personal_info(doc)
    header_candidates = []
    content_list = list(iterate_doc_content(doc))
    
    for i, (elem_type, element) in enumerate(content_list):
        sect_type, score = is_section_header(element)
        if sect_type != SectionType.UNKNOWN:
            # BOOST: Prximity to placeholder
            has_placeholder = False
            for look in range(1, 10):
                if i + look < len(content_list):
                    nxt = content_list[i+look][1].text.upper()
                    if "[FILL" in nxt or "{{" in nxt: has_placeholder = True; break
            
            final_confidence = score + (0.3 if has_placeholder else 0.0)
            header_candidates.append({
                'index': i, 'type': sect_type, 'confidence': final_confidence,
                'text': element.text.strip(), 'elem_type': elem_type,
                'table_idx': getattr(element, '_parent_table_idx', None), 'xml_id': id(element._element)
            })

    candidates = personal_candidates + header_candidates
    clean_candidates = _denoise_candidates(candidates)
    sections = []
    content_list = list(iterate_doc_content(doc))
    for i, cand in enumerate(clean_candidates):
        _, element = content_list[cand['index']]
        mark_protected_xml(element)
        sections.append(TemplateSection(
            section_index=i, section_type=cand['type'], confidence=cand['confidence'],
            start_paragraph_idx=cand['index'] if cand['elem_type'] == 'para' else None,
            table_idx=cand['table_idx'], header_text=cand['text'], raw_heading=cand['text'],
            header_xml_id=cand['xml_id'], header_element_type=cand['elem_type']
        ))
        logger.info(f"Final Section [{i}]: {cand['type'].name} - '{cand['text']}' (MARKED PROTECTED)")

    return TemplateSchema(template_name=template_name, template_file=template_path, sections=sections)
