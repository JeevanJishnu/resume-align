from docx import Document
from docx.shared import Pt

template_path = r'templates\AM0275 Vishnuraj K J.docx'
doc = Document(template_path)
table = doc.tables[0]

# Force Name Box
p = table.rows[0].cells[0].paragraphs[0]
p.text = "[fill Name here]"
p.runs[0].bold = True
p.runs[0].font.size = Pt(20)

# Clear ghost lines in other cells
for r in range(1, len(table.rows)):
    for c in range(len(table.rows[r].cells)):
        cell = table.rows[r].cells[c]
        # Just ensure headers are bold and placeholders exist
        pass

doc.save(template_path)
print("Template Fixed: [fill Name here] in T0R0")
