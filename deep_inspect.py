from docx import Document
from template_engine.template_cleaner import is_protected_xml

doc = Document('templates/AM0275 Vishnuraj K J.docx')
print("--- AM0275 STRUCTURE ---")
for t_idx, t in enumerate(doc.tables):
    for r_idx, r in enumerate(t.rows):
        for c_idx, c in enumerate(r.cells):
            print(f"T{t_idx} R{r_idx} C{c_idx}:")
            for p_idx, p in enumerate(c.paragraphs):
                print(f"  P{p_idx}: [{p.text.strip()}]")
            for nt_idx, nt in enumerate(c.tables):
                print(f"  Nested Table {nt_idx}")
