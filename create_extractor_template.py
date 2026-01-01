
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_extractor_template(output_path):
    doc = Document()
    
    # Title Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # HEADER
    h1 = doc.add_paragraph()
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = h1.add_run("[NAME]")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_contact.add_run("Email: {{email}} | Phone: {{phone}} | LinkedIn: {{linkedin}}")

    doc.add_paragraph("_" * 50).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 1. SUMMARY
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph("[FILL HERE]")

    # 2. SKILLS
    doc.add_heading('TECHNICAL SKILLS', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Skills / Tools'
    
    # Dynamic row for mapping
    row_cells = table.add_row().cells
    row_cells[0].text = "Category" 
    row_cells[1].text = "[FILL HERE]"

    # 3. WORK EXPERIENCE
    doc.add_heading('WORK EXPERIENCE', level=1)
    
    # Create the "Blueprint" block for experience
    # We use a table for better copy-paste layout
    exp_table = doc.add_table(rows=2, cols=2)
    exp_table.style = 'Table Grid'
    
    # Row 1: Role & Company
    r1 = exp_table.rows[0].cells
    p1 = r1[0].paragraphs[0]
    p1.add_run("Role: ").bold = True
    p1.add_run("[FILL HERE]")
    
    p2 = r1[1].paragraphs[0]
    p2.add_run("Company: ").bold = True
    p2.add_run("[FILL HERE]")

    # Row 2: Duration & Description
    r2 = exp_table.rows[1].cells
    # Merge for full width description
    r2[0].merge(r2[1])
    
    p3 = r2[0].paragraphs[0]
    p3.add_run("Duration: ").bold = True
    p3.add_run("[FILL HERE]")
    p3.add_run("\n")
    p3.add_run("Responsibilities:").bold = True
    p3.add_run("\n[FILL HERE]")
    
    doc.add_paragraph("") # Spacer

    # 4. PROJECTS
    doc.add_heading('PROJECTS', level=1)
    
    # Blueprint for Projects
    proj_table = doc.add_table(rows=4, cols=2)
    proj_table.style = 'Table Grid'
    
    # Project Title
    proj_table.rows[0].cells[0].text = "Project Title:"
    proj_table.rows[0].cells[1].text = "[FILL HERE]"
    
    # Tech Stack
    proj_table.rows[1].cells[0].text = "Technologies:"
    proj_table.rows[1].cells[1].text = "[FILL HERE]"
    
    # Role
    proj_table.rows[2].cells[0].text = "Role:"
    proj_table.rows[2].cells[1].text = "[FILL HERE]"
    
    # Description
    proj_table.rows[3].cells[0].text = "Description:"
    proj_table.rows[3].cells[1].text = "[FILL HERE]"
    
    doc.add_paragraph("")

    # 5. EDUCATION
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph("• [FILL HERE]")

    # 6. CERTIFICATIONS
    doc.add_heading('CERTIFICATIONS', level=1)
    doc.add_paragraph("• [FILL HERE]")

    doc.save(output_path)
    print(f"Created Extractor Template at {output_path}")

if __name__ == "__main__":
    os.makedirs("templates", exist_ok=True)
    create_extractor_template("templates/Extractor_Master.docx")
