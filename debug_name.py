import re
from docx import Document
from fuzzywuzzy import fuzz

def norm(t):
    return re.sub(r'[\[\]\(\):]', '', t).strip().lower()

targets = [
    ("Corp_Dev", "templates/Corp_Dev.docx"),
    ("AM0275", "templates/AM0275 Vishnuraj K J.docx")
]

oh = norm("[fill Name here]")
print(f"Target Norm: '{oh}'")

for name, path in targets:
    print(f"\n--- Checking {name} ---")
    doc = Document(path)
    for t_idx, t in enumerate(doc.tables):
        for r_idx, r in enumerate(t.rows):
            for c_idx, c in enumerate(r.cells):
                for p_idx, p in enumerate(c.paragraphs):
                    n = norm(p.text)
                    ratio = fuzz.ratio(n, oh)
                    if ratio > 50 or "fill" in n:
                        print(f"T{t_idx} R{r_idx} C{c_idx} P{p_idx}: Text='{p.text}' Norm='{n}' Ratio={ratio}")
