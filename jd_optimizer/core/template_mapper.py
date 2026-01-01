import re
import os
import logging
from typing import List, Dict, Any
from copy import deepcopy
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from .template_models import SectionType, TemplateSchema
from .template_extractor import iterate_doc_content, normalize_text, is_section_header, SECTION_SYNONYMS

logger = logging.getLogger("TemplateMapper")
logger.setLevel(logging.DEBUG)

def _fuzzy_compare(s1, s2):
    from fuzzywuzzy import fuzz
    if not s1 or not s2: return False
    return fuzz.ratio(str(s1).lower(), str(s2).lower()) > 85

def fill_template(schema: TemplateSchema, data: Dict[str, Any], output_path: str):
    """
    Fills a cleaned template with CV data.
    Uses dynamic anchor detection to be resilient to document shifts.
    """
    try:
        doc = Document(schema.template_file)
        
        # 1. Global Replacements (Personal Info)
        personal_map = {
            "[fill Name here]": data.get("full_name", "Applicant"),
            "[NAME]": data.get("full_name", "Applicant"),
            "{{name}}": data.get("full_name", "Applicant"),
            "{{email}}": data.get("email", ""),
            "{{phone}}": data.get("phone", ""),
            "{{linkedin}}": data.get("linkedin", "")
        }
        
        for p in doc.paragraphs:
            for k, v in personal_map.items():
                if k.lower() in p.text.lower():
                    # Preserve formatting if possible
                    p.text = p.text.replace(k, v)
        
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for k, v in personal_map.items():
                            if k.lower() in p.text.lower():
                                p.text = p.text.replace(k, v)

        # 2. Block Replication (Dynamic Scaling)
        try:
            _replicate_dynamic_sections(doc, data)
        except Exception as e:
            logger.error(f"Dynamic replication failed (skipping): {e}")

        # 3. Sequential Scanning
        content_list = list(iterate_doc_content(doc))
        from .template_extractor import is_section_header
        
        # Detect all section anchors in the LIVE document
        anchors = [] # List of (index, section_type)
        for i, (etype, elem) in enumerate(content_list):
            if etype == 'table': continue
            stype, score, is_record = is_section_header(elem)
            if stype != SectionType.UNKNOWN and score >= 0.7:
                # Distinguish top-level vs record-level
                is_top_level = not is_record
                
                # If it's a top-level section type, or the first time we see this type
                if is_top_level or not anchors or anchors[-1][1] != stype:
                    anchors.append((i, stype))
        
        logger.info(f"Detected {len(anchors)} anchors in document for mapping: {[a[1].name for a in anchors]}")
        
        # Skill Categories for smart mapping
        SKILL_CATEGORIES = {
            "cloud": ["aws", "azure", "gcp", "lambda", "ec2", "s3", "cloud"],
            "devops": ["docker", "kubernetes", "jenkins", "terraform", "ansible", "git", "github", "gitlab"],
            "database": ["sql", "mysql", "postgres", "mongodb", "redis", "oracle", "elasticsearch"],
            "os": ["linux", "windows", "unix", "ubuntu"],
            "web": ["html", "css", "javascript", "react", "angular", "vue", "node", "express"],
            "language": ["python", "java", "php", "c#", "ruby", "typescript"],
            "design": ["figma", "photoshop", "illustrator", "sketch", "xd", "ui", "ux"]
        }
        
        allocated_skills = set()
        section_processed = set() # Track unique anchor indices
        section_record_index = {} # SectionType -> Current Record Index

        # Loop through each detected anchor
        last_found_idx = -1
        for a_idx, (start_idx, stype) in enumerate(anchors):
            # Define target range (from this anchor to the next anchor)
            end_idx = anchors[a_idx+1][0] if a_idx + 1 < len(anchors) else len(content_list)
            
            # Collect targets in this range
            targets = []
            for i in range(start_idx + 1, end_idx):
                etype, elem = content_list[i]
                if etype == 'table': continue
                txt = elem.text.strip()
                if "[FILL HERE]" in txt or "[fill" in txt.lower():
                    targets.append((i, etype, elem))
            
            if not targets: continue
            
            # Fetch relevant data for this section type
            list_data = [] # List of dicts or list of strings
            if stype == SectionType.SUMMARY: list_data = [data.get("summary", "")]
            elif stype == SectionType.SKILLS or stype == SectionType.TOOLS:
                s_list = data.get("skills", [])
                t_list = data.get("tools", [])
                list_data = list(dict.fromkeys(s_list + t_list))
            elif stype == SectionType.WORK_EXPERIENCE: list_data = data.get("work_experience", [])
            elif stype == SectionType.PROJECTS: list_data = data.get("projects", [])
            elif stype == SectionType.EDUCATION: list_data = data.get("education", [])
            elif stype == SectionType.CERTIFICATIONS: list_data = data.get("certifications", [])

            if not list_data:
                # Clear all targets if no data
                for _, _, elem in targets: elem.clear()
                continue

            # Fill Targets
            is_skill_section = (stype in [SectionType.SKILLS, SectionType.TOOLS])
            
            if stype not in section_record_index:
                section_record_index[stype] = {"rec": 0, "field": 0, "last_block": -1}
            
            state = section_record_index[stype]
            is_sequential = (stype in [SectionType.WORK_EXPERIENCE, SectionType.PROJECTS])
            
            for i, (full_idx, etype, elem) in enumerate(targets):
                # ADVANCE RECORD?
                if is_sequential:
                    look_start = targets[i-1][0] if i > 0 else start_idx
                    for check_idx in range(look_start + 1, full_idx):
                        c_etype, c_elem = content_list[check_idx]
                        if c_etype != 'table':
                            c_txt = c_elem.text.strip()
                            # Match "#N" or "Project #N"
                            match = re.search(r'#(\d+)', c_txt)
                            if match:
                                num = int(match.group(1))
                                if state['last_block'] != check_idx:
                                    state['rec'] = num - 1
                                    state['field'] = 0
                                    state['last_block'] = check_idx
                                    logger.debug(f"Advancing record to index {state['rec']} at '{c_txt}'")
                            elif "#" in c_txt or (len(c_txt) < 30 and re.match(r"^(project|assignment|task)\s*#?\d*$", c_txt.lower())):
                                if state['last_block'] != check_idx:
                                    state['rec'] += 1
                                    state['field'] = 0
                                    state['last_block'] = check_idx
                                    logger.debug(f"Advancing record to index {state['rec']} at '{c_txt}'")

                # Identify Label
                label_text = ""
                table = getattr(elem, '_parent_table', None)
                t_idx = id(table) if table else -1
                
                # Check paragraph above for label
                if full_idx > start_idx + 1:
                    prev_etype, prev_elem = content_list[full_idx - 1]
                    # Table objects don't have .text attribute
                    if prev_etype != 'table' and hasattr(prev_elem, 'text'):
                        if prev_elem.text.strip().endswith(':') or len(prev_elem.text.strip()) < 30:
                            label_text = prev_elem.text.strip().lower().rstrip(':').strip()
                
                # Check Col 0 of table for label
                if not label_text and etype == "table_cell" and table:
                    target_id = id(elem._element)
                    for row in table.rows:
                        if any(id(p._element) == target_id for cell in row.cells for p in cell.paragraphs):
                            label_text = row.cells[0].text.strip().lower().rstrip(':').strip()
                            break
                            
                # Check element for label in its own text (e.g. "Title: [FILL HERE]")
                if not label_text:
                    if ":" in elem.text and "[" in elem.text and elem.text.find(":") < elem.text.find("["):
                        label_text = elem.text.split(":", 1)[0].strip().lower()
                
                # Record Advancement (Sequential for non-skills)
                if not is_skill_section and state["field"] > 0:
                    is_new_block = (t_idx != -1 and state["last_block"] != -1 and t_idx != state["last_block"])
                    # More specific repeat check: only increment if it's a primary identification field
                    is_repeat_title = False
                    if label_text:
                        lt_clean = label_text.lower()
                        # MUST contain title/name/employer OR project/company as a key word
                        if any(k in lt_clean for k in ["title", "project name", "company", "employer"]):
                            is_repeat_title = True
                        elif "project" in lt_clean and ("title" in lt_clean or "name" in lt_clean):
                            is_repeat_title = True
                            
                    if is_new_block or is_repeat_title:
                        state["rec"] += 1; state["field"] = 0
                        logger.info(f"Advancing record for {stype} to index {state['rec']} (Label: {label_text})")
                
                state["last_block"] = t_idx
                val = None
                
                if is_skill_section:
                    # Skill Filtering Logic
                    SKILL_CATEGORIES = {
                        "category": ["category", "key skills"],
                        "tools": ["tools", "technologies", "software"],
                        "environment": ["environment", "tech stack", "stack"]
                    }
                    if label_text and any(k in label_text for k in ["category", "tools", "technologies", "environment"]):
                        matches = []
                        for cat_key, keywords in SKILL_CATEGORIES.items():
                            if cat_key in label_text:
                                for s in list_data:
                                    if s not in allocated_skills and any(k in s.lower() for k in keywords):
                                        matches.append(s)
                                        allocated_skills.add(s)
                        if matches: val = ", ".join(matches)
                    else:
                        # Generic skill list for the section
                        remaining = [s for s in list_data if s not in allocated_skills]
                        if remaining and start_idx not in section_processed:
                            has_table = any(t[1] == "table_cell" for t in targets)
                            if etype == "table_cell" or not has_table or i == len(targets) - 1:
                                val = ", ".join(remaining)
                                allocated_skills.update(remaining)
                                section_processed.add(start_idx)
                else:
                    # Complex Record Mapping
                    if state["rec"] >= len(list_data):
                        # Empty data: If it's a paragraph, remove it to avoid empty bullets
                        if etype == "para":
                            try:
                                p_elem = elem._element
                                p_elem.getparent().remove(p_elem)
                            except:
                                elem.clear()
                        else:
                            elem.clear()
                        continue
                        
                    rec = list_data[state["rec"]]
                    if isinstance(rec, str):
                        val = rec
                    else:
                        lt = label_text
                        if not lt: val = None
                        elif any(k in lt for k in ["role", "designation", "position"]): val = rec.get("role", "")
                        elif any(k in lt for k in ["duration", "period", "date", "year"]): val = rec.get("duration", "")
                        elif "client" in lt: val = rec.get("client", "")
                        elif any(k in lt for k in ["title", "project", "name", "company", "employer"]): 
                            val = rec.get("title", rec.get("project_name", rec.get("company", "")))
                        elif any(k in lt for k in ["tech", "stack", "env", "tool", "technologies"]): 
                            val = rec.get("tech", rec.get("tech_stack", rec.get("environment", "")))
                        elif any(k in lt for k in ["resp", "desc", "detail", "overview", "task"]):
                            # Try both "details" and "description"
                            val = rec.get("details", rec.get("responsibilities", rec.get("description", "")))
                        elif any(k in lt for k in ["inst", "univ", "school"]): val = rec.get("institution", "")
                        elif any(k in lt for k in ["deg", "qual"]): val = rec.get("degree", "")

                    if val is None and state["field"] == 0:
                        has_table_targets = any(t[1] == "table_cell" for t in targets)
                        if not is_sequential or not has_table_targets:
                            if start_idx not in section_processed:
                                _fill_list_to_paragraph(elem, stype, list_data)
                                val = "DONE"; section_processed.add(start_idx)

                # Inject Value
                if val == "DONE": 
                    val = None
                    state["field"] += 1 
                
                if val:
                    logger.debug(f"Injecting into {stype.name} rec {state['rec']} field {state['field']}: '{label_text}' -> '{str(val)[:50]}'")
                    # If label was in-place, preserve it
                    if ":" in elem.text and "[" in elem.text:
                        label_part = elem.text.split("[", 1)[0]
                        elem.clear()
                        elem.add_run(label_part + str(val))
                    else:
                        elem.clear()
                        elem.add_run(str(val))
                    state["field"] += 1
                    
                    if not is_sequential and not is_skill_section:
                        state["rec"] += 1
                        state["field"] = 0
                else:
                    # If it's a bulleted paragraph, remove it if empty
                    if etype == "para":
                        try:
                            p_elem = elem._element
                            p_elem.getparent().remove(p_elem)
                        except:
                            elem.clear()
                    else:
                        elem.clear()

        doc.save(output_path)
        logger.info(f"SUCCESS: Saved document to {output_path}")
    except Exception as e:
        logger.exception(f"Error in fill_template: {e}")

def _fill_list_to_paragraph(target_p, section_type, items):
    for item in reversed(items):
        if not item: continue
        p = target_p.insert_paragraph_before()
        p.paragraph_format.space_after = 0
        if section_type == SectionType.WORK_EXPERIENCE:
            role = item.get('role', 'Role')
            comp = item.get('company', 'Company')
            p.add_run(f"• {role} at {comp} ({item.get('duration', 'N/A')})").bold = True
            desc = item.get('responsibilities', '')
            if desc: 
                for d_line in reversed(str(desc).split('\n')):
                    if d_line.strip():
                        dp = target_p.insert_paragraph_before(d_line.strip())
                        dp.paragraph_format.space_after = 0
        elif section_type == SectionType.PROJECTS:
            title = item.get('title', item.get('project_name', 'Project'))
            p.add_run(f"• {title} ({item.get('duration', 'N/A')})").bold = True
        elif section_type == SectionType.EDUCATION:
            deg = item.get('degree', 'Degree')
            inst = item.get('institution', 'N/A')
            p.add_run(f"• {deg} from {inst} ({item.get('duration', 'N/A')})").bold = True
        elif section_type == SectionType.CERTIFICATIONS:
            title = item if isinstance(item, str) else item.get('title', 'Certification')
            p.add_run(f"• {title}")
        elif section_type == SectionType.SUMMARY:
            p.add_run(str(item))
        else:
            p.add_run(f"• {str(item)}")
    target_p.clear()
    
def _sanitize_xml(element):
    """Removes Word-specific IDs that cause 'unreadable content' errors when duplicated."""
    # Namespaces for paraId and textId
    w14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
    for e in element.iter():
        for attr in [f"{w14}paraId", f"{w14}textId"]:
            if attr in e.attrib:
                del e.attrib[attr]

def _replicate_dynamic_sections(doc, data):
    """
    Identifies 'blueprint' blocks in repeatable sections and clones them
    to match the number of records in the data.
    """
    SECTION_RECORDS = {
        SectionType.PROJECTS: data.get("projects", []),
        SectionType.WORK_EXPERIENCE: data.get("work_experience", []),
        SectionType.EDUCATION: data.get("education", [])
    }
    
    # 1. Get Top Level elements (not nested in tables)
    top_elements = []
    # Use doc.element.body.iterchildren() to get direct children
    for child in doc.element.body.iterchildren():
        tag = getattr(child, 'tag', '')
        if tag.endswith('}p') or tag == 'p':
            top_elements.append(Paragraph(child, doc))
        elif tag.endswith('}tbl') or tag == 'tbl':
            top_elements.append(Table(child, doc))

    if not top_elements: 
        logger.warning("No top-level elements found in document body.")
        return

    # 2. Assign sections to elements
    element_sections = []
    current_section = SectionType.UNKNOWN
    
    for elem in top_elements:
        stype, score, is_record = is_section_header(elem)
        
        # Check if the current element is a top-level section header
        is_top_level_section_header = (stype != SectionType.UNKNOWN and score >= 0.7 and not is_record)

        if is_top_level_section_header:
            if stype != current_section:
                logger.info(f"Template Section transition: {current_section.name} -> {stype.name} at '{elem.text.strip()}'")
            current_section = stype
            element_sections.append({'elem': elem, 'section': current_section, 'is_section_title': True})
        else:
            element_sections.append({'elem': elem, 'section': current_section, 'is_section_title': False})

    # 3. Process Repeatable Sections
    for stype, records in SECTION_RECORDS.items():
        section_elements = [e for e in element_sections if e['section'] == stype]
        if not section_elements: continue
        
        # Group elements into "Blocks"
        # A "Block" is a logical group of elements representing one record (e.g. A Project Title + Description Table)
        blocks = [] 
        curr_block = []
        
        for e in section_elements:
            elem = e['elem']
            is_header = e['is_section_title']
            
            # If we hit a Section Header (e.g. "PROJECTS"), it ALWAYS starts a new block
            if is_header:
                if curr_block: blocks.append(curr_block)
                curr_block = [elem]
            else:
                # For content, we group it. 
                # If we detect a "Record Header" (like Project #1), that also forces a new block start.
                _, _, is_record_start = is_section_header(elem)
                if is_record_start and curr_block: 
                    # But wait, if the current block is JUST the section header, we add to it? 
                    # No, usually Top Header is separate.
                    blocks.append(curr_block)
                    curr_block = [elem]
                elif not curr_block:
                    curr_block = [elem]
                else:
                    curr_block.append(elem)
                    
        if curr_block: blocks.append(curr_block)
        
        # --- CRITICAL FIX: Header Separation ---
        # The first block might contain the Section Header ("Projects").
        # We MUST NOT replicate the header.
        # We need to find the "Blueprint" block which is the CONTENT structure.
        
        blueprint_blocks = []
        
        if blocks:
            first_block = blocks[0]
            # Check if first element of first block is the Section Header
            first_elem_wrapper = next((x for x in section_elements if x['elem'] == first_block[0]), None)
            
            if first_elem_wrapper and first_elem_wrapper['is_section_title']:
                # The first block STARTS with the header.
                # If the block has MORE than just the header (e.g. Header + Table), we split it.
                if len(first_block) > 1:
                    # Header is index 0. Content is 1..N
                    # We leave index 0 alone (static). We take 1..N as a candidate for blueprint.
                    blueprint_blocks.append(first_block[1:])
                    # Add any subsequent blocks as is
                    blueprint_blocks.extend(blocks[1:])
                else:
                    # The first block is ONLY the header. We ignore it for replication.
                    blueprint_blocks.extend(blocks[1:])
            else:
                # First block is not a header? Then it's content.
                blueprint_blocks.extend(blocks)
                
        if not blueprint_blocks:
            logger.warning(f"No content blocks found for {stype} to replicate.")
            continue
            
        # The "Blueprint" is the LAST block found (assuming it represents a single empty record structure)
        # Note: In an Extractor template, we usually have Header + 1 Empty Record Block.
        # So blueprint_blocks should have length 1.
        
        blueprint = blueprint_blocks[-1]
        
        num_existing = len(blueprint_blocks)
        num_needed = len(records)
        
        logger.info(f"Replication {stype}: Found {num_existing} content blocks. Need {num_needed}.")

        if num_needed > num_existing:
            # We need to Clone the Blueprint (N - Existing) times
            # We append clones AFTER the last element of the last block
            last_elem_in_doc = blueprint_blocks[-1][-1]
            
            for i in range(num_existing, num_needed):
                # Clone the entire blueprint block
                new_block = []
                for b_elem in blueprint:
                    new_xml = deepcopy(b_elem._element)
                    _sanitize_xml(new_xml)
                    last_elem_in_doc._element.addnext(new_xml)
                    
                    if isinstance(b_elem, Paragraph):
                        new_obj = Paragraph(new_xml, doc)
                        # Increment number if present
                        if "#" in new_obj.text:
                            new_obj.text = re.sub(r'#\d+', f'#{i+1}', new_obj.text)
                    else:
                        new_obj = Table(new_xml, doc)
                    
                    new_block.append(new_obj)
                    last_elem_in_doc = new_obj # Update cursor
                    
            logger.info(f"Replicated {stype}: Scaled up to {num_needed} blocks.")
            
        elif num_needed < num_existing and num_needed > 0:
            # Scale down (optional, but good for cleanup)
            # Remove excess blocks from the end
            for i in range(num_existing - 1, num_needed - 1, -1):
                block_to_remove = blueprint_blocks[i]
                for b_elem in block_to_remove:
                    try:
                        b_elem._element.getparent().remove(b_elem._element)
                    except:
                        pass
            logger.info(f"Replicated {stype}: Scaled down to {num_needed} blocks.")
    
    logger.info("Dynamic block replication complete.")
