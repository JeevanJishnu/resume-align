
import re
import logging
from docx import Document
from fuzzywuzzy import fuzz
from .template_models import TemplateSchema, SectionType
from .template_extractor import iterate_doc_content, is_section_header

logger = logging.getLogger("NuclearCleaner")
logger.setLevel(logging.INFO)

def clean_template_content(doc: Document, schema: TemplateSchema) -> Document:
    """
    NUCLEAR CLEANER v5.1:
    - Sequentially identifies Headers, Labels, and Data.
    - Preserves Headers and Labels.
    - Scrubs Data and replaces with [FILL HERE].
    """
    content_list = list(iterate_doc_content(doc))
    to_delete_paras = []
    
    # Pre-scan for headers
    header_found_at = {}
    first_header_idx = 9999
    for i, (etype, elem) in enumerate(content_list):
        if etype == 'table': continue
        stype, score, _ = is_section_header(elem)
        if stype != SectionType.UNKNOWN and score >= 0.7:
            header_found_at[i] = stype
            first_header_idx = min(first_header_idx, i)
            
    PROTECTED_LABELS = [
        "client", "role", "duration", "company", "location", "responsibilities", 
        "description", "details", "technologies", "tech", "stack", "environment",
        "institution", "degree", "qualification", "project", "title", "summary",
        "tools", "category", "skills", "experience", "education", "overview"
    ]

    filled_cells = set()

    for i, (etype, elem) in enumerate(content_list):
        if etype == 'table': continue
        txt = elem.text.strip()
        
        # A. Section Header
        if i in header_found_at:
            mtype = header_found_at[i]
            if mtype == SectionType.FULL_NAME:
                elem.clear()
                elem.add_run("[fill Name here]")
            continue

        # B. Global Placeholder protection
        if any(x in txt for x in ["{{", "[fill", "[FILL", "[NAME]"]):
            continue

        # C. Label Detection
        table_obj = getattr(elem, '_parent_table', None)
        is_label = False
        t_clean = txt.lower().rstrip(':').strip()
        
        if txt.endswith(':') and len(txt.split()) < 6:
            is_label = True
        elif ":" in txt and len(txt.split()) < 5:
            if len(txt.split(':', 1)[1].strip()) < 3:
                is_label = True
        elif table_obj and any(txt == r.cells[0].text.strip() for r in table_obj.rows):
            if t_clean in PROTECTED_LABELS or (len(txt.split()) < 4 and t_clean):
                is_label = True
        
        if is_label:
            continue

        # D. Data Scrubbing
        if not txt: continue
        
        if i < first_header_idx:
            # Top Matter
            words = [w for w in txt.split() if w.isalpha()]
            is_name_like = (len(words) >= 2 and len(words) <= 5 and (txt.isupper() or all(w[0].isupper() for w in words)))
            is_contact = "@" in txt or "linkedin" in txt or "http" in txt or re.search(r'\+?\d[\d\-\s]{7,}', txt)
            
            if is_name_like:
                elem.clear()
                elem.add_run("[fill Name here]")
            elif is_contact or len(txt) < 100:
                elem.clear()
            else:
                elem.clear()
                elem.add_run("[FILL HERE]")
        else:
            # Body content
            if etype == "table_cell":
                cell_id = id(elem._element.getparent())
                if cell_id not in filled_cells:
                    elem.clear()
                    elem.add_run("[FILL HERE]")
                    filled_cells.add(cell_id)
                else:
                    to_delete_paras.append(elem)
            else:
                elem.clear()
                elem.add_run("[FILL HERE]")

    for p in to_delete_paras:
        try:
            p._element.getparent().remove(p._element)
        except:
            p.clear()

    # E. Ensure all non-label cells have at least one [FILL HERE]
    for etype, elem in content_list:
        if etype == 'table':
            for row in elem.rows:
                first_cell_txt = row.cells[0].text.strip().lower().rstrip(':')
                is_row_labeled = any(fuzz.partial_ratio(first_cell_txt, lbl) > 85 for lbl in PROTECTED_LABELS) or row.cells[0].text.strip().endswith(':')
                
                for idx, cell in enumerate(row.cells):
                    if idx == 0 and is_row_labeled: continue
                    ctxt = cell.text.strip()
                    if not any(x in ctxt for x in ["{{", "[fill", "[FILL", "[NAME]"]):
                        if not cell.paragraphs:
                            cell.add_paragraph("[FILL HERE]")
                        else:
                            cell.paragraphs[0].text = "[FILL HERE]"
                            for extra_p in cell.paragraphs[1:]:
                                try: extra_p._element.getparent().remove(extra_p._element)
                                except: pass

    return doc
